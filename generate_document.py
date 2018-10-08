from docx import Document
from win32com import client
import os
wdFormatPDF = 17
def covx_to_pdf(infile, outfile):
    """Convert a Word .docx to PDF"""
    word = client.DispatchEx('Word.Application')
    doc = word.Documents.Open(infile)
    doc.SaveAs(outfile, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    os.remove(infile)

def replace_in_paragraphs(paragraphs, src, dest):
    dest = str(dest)
    if (dest == 'None'):
        dest = ''

    for p in paragraphs:
        if src in p.text:
            inline = p.runs
            str_concate = ''
            for i in range(len(inline)):
                str_concate += inline[i].text
            pos_src = str_concate.find(src)
            if (pos_src >= 0):
                str_idx = 0
                replace_flag = False
                for i in range(len(inline)):
                    len_text = len(inline[i].text)
                    if (str_idx <= pos_src and pos_src < str_idx + len_text):
                        if (replace_flag == False):
                            inline[i].text = inline[i].text[:pos_src - str_idx] + dest + inline[i].text[pos_src - str_idx + len(src):]
                            replace_flag = True
                    elif (replace_flag == True and pos_src <= str_idx and str_idx < pos_src + len(src)):
                        inline[i].text = ''
                    str_idx += len_text

def replace_docx(src_file, dest_file, replace_list):
    if (dest_file[-4] == '.pdf'):
        return -1 #the output file must be PDF

    doc = Document(src_file)

    for i in range(len(replace_list)):
        replace_in_paragraphs(doc.paragraphs, replace_list[i][0], replace_list[i][1])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for i in range(len(replace_list)):
                    replace_in_paragraphs(cell.paragraphs, replace_list[i][0], replace_list[i][1])

    doc.save(dest_file[0:-3] + 'docx')
    covx_to_pdf(dest_file[0:-3] + 'docx', dest_file + '.pdf')


